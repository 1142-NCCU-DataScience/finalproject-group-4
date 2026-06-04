"""
build_word_report.py
從 baseball_project/ 執行：
    python scripts/build_word_report.py
輸出：docs/final_report.docx
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import docx.opc.constants

ROOT    = Path(__file__).resolve().parent.parent
FIG_DIR = ROOT / "figures"
POST_DIR = FIG_DIR / "poster"
OUT     = ROOT / "docs" / "final_report.docx"

# ── 輔助函數 ──────────────────────────────────────────────
def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def add_para(doc, text, bold=False, italic=False, size=11, color=None, indent=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if indent:
        p.paragraph_format.left_indent = Cm(1)
    return p

def add_figure(doc, img_path, caption, badge=None, width_cm=15):
    """插入圖片 + 說明文字 + 海報/簡報標籤"""
    if not Path(img_path).exists():
        doc.add_paragraph(f"[圖片未找到: {img_path}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(img_path), width=Cm(width_cm))

    # badge 行
    if badge:
        bp = doc.add_paragraph()
        bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        br = bp.add_run(badge)
        br.bold = True
        br.font.size = Pt(9)
        if "海報" in badge:
            br.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        else:
            br.font.color.rgb = RGBColor(0x27, 0x6F, 0xBF)

    # caption
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cp.add_run(caption)
    cr.italic = True
    cr.font.size = Pt(9)
    cr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    doc.add_paragraph()   # 間距

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    # 標頭
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 標頭背景色
        tc = hdr[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "2E75B6")
        tcPr.append(shd)
        for run in hdr[i].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # 內容
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cells[ci].paragraphs[0].runs:
                run.font.size = Pt(9.5)
    if col_widths:
        for ri in range(len(table.rows)):
            for ci, w in enumerate(col_widths):
                table.rows[ri].cells[ci].width = Cm(w)
    doc.add_paragraph()

def add_warning_box(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.8)
    p.paragraph_format.right_indent = Cm(0.8)
    shading_elm = OxmlElement("w:pPr")
    run = p.add_run("⚠️  " + text)
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x7B, 0x34, 0x1E)
    run.bold = True


# ── 建立文件 ─────────────────────────────────────────────
doc = Document()

# 全域字體設定（中文 + 英文）
style = doc.styles["Normal"]
style.font.name = "Arial"
style.font.size = Pt(11)
# 中文字體
from docx.oxml.ns import qn
style.element.rPr.rFonts.set(qn("w:eastAsia"), "微軟正黑體")

# 頁面邊距
from docx.oxml import OxmlElement
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

# ══════════════════════════════════════════════════════════
# 封面
# ══════════════════════════════════════════════════════════
doc.add_paragraph()
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title_p.add_run("MLB 打者生涯預測分析報告")
tr.bold = True
tr.font.size = Pt(22)
tr.font.color.rgb = RGBColor(0x1F, 0x45, 0x7E)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub_p.add_run("Player Career Projection — 資料科學期末專題")
sr.font.size = Pt(14)
sr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
sr.italic = True

doc.add_paragraph()
for line in [
    "作者：Brian 吳帛恩",
    "課程：資料科學（2026 春季學期）",
    "資料時間範圍：2018–2024 MLB 球季",
    "報告日期：2026-05-16",
]:
    lp = doc.add_paragraph()
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lr = lp.add_run(line)
    lr.font.size = Pt(11)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# §1 研究背景與目標
# ══════════════════════════════════════════════════════════
add_heading(doc, "1. 研究背景與目標", 1)
add_para(doc,
    "傳統棒球評估長期仰賴打擊率（AVG）、打點（RBI）等統計，但這些指標容易受球場、"
    "隊友、運氣因素扭曲，無法真實反映球員個人能力。現代賽伯計量學（Sabermetrics）"
    "引入 wRC+（加權上壘打點創造值，聯盟平均 = 100）和 WAR（勝場貢獻值）等進階指標，"
    "提供更準確的打者表現衡量。")

add_para(doc, "核心問題：", bold=True)
add_para(doc,
    "能否僅用球員過去三個球季的「打擊過程指標」，預測其下一季的 wRC+？",
    italic=True, indent=True)

add_para(doc, "研究目標：", bold=True)
for item in [
    "建立以 Statcast 為唯一資料來源的打擊指標計算管線",
    "透過機器學習（Random Forest）預測球員隔年 wRC+",
    "使用嚴格的時間切分（依目標球季）避免未來資料洩漏",
    "分析哪些過程指標對預測最具貢獻",
]:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(item)
    run.font.size = Pt(11)

# ══════════════════════════════════════════════════════════
# §2 資料來源與收集方式
# ══════════════════════════════════════════════════════════
add_heading(doc, "2. 資料來源與收集方式", 1)

add_heading(doc, "2.1 主要資料來源", 2)
add_table(doc,
    headers=["資料來源", "取得方式", "用途", "時間範圍"],
    rows=[
        ["Baseball Savant Statcast", "pybaseball.statcast()", "逐球事件、計算所有打擊率指標", "2018–2024"],
        ["Lahman Database", "pybaseball.lahman.fielding()", "球員守備位置（primary position）", "2018–2024"],
        ["Chadwick Register", "pybaseball.playerid_reverse_lookup()", "MLBAM ID → 球員姓名對應", "輔助查詢"],
    ],
    col_widths=[4.5, 4.5, 5.5, 3]
)

add_para(doc,
    "重要設計決策：本專案完全不使用 FanGraphs 等外部打擊排行榜。所有比率指標（K%、BB%、"
    "BABIP、BIP%、wRC+）均由 Statcast 逐球事件資料自行計算，確保資料來源一致性與可重現性。",
    italic=True)

add_heading(doc, "2.2 Statcast 資料結構", 2)
add_para(doc,
    "Statcast 為逐球（pitch-by-pitch）事件資料，每一列代表一個投球事件。"
    "本專案保留打席結束列（events 欄位非空），每個球季約 70–90 萬筆資料，"
    "存為 .parquet 格式（約 100–200 MB/季），共下載 2018–2024 年共 7 個球季。")

add_table(doc,
    headers=["欄位名稱", "說明"],
    rows=[
        ["batter", "打者 MLBAM ID"],
        ["game_year", "球季年份"],
        ["events", "打席結果（single, home_run, strikeout 等 22 種）"],
        ["plate_x / plate_z", "進壘水平/垂直位置（ft，捕手視角）"],
        ["woba_value / woba_denom", "wOBA 計算用數值"],
        ["age_bat", "打者當時年齡"],
        ["launch_speed / launch_angle", "擊球初速（mph）與仰角（度）"],
        ["estimated_ba_using_speedangle", "期望打擊率 xBA（Statcast 物理模型）"],
    ],
    col_widths=[5.5, 12]
)

add_heading(doc, "2.3 特定球員資料（熱區圖用）", 2)
add_para(doc,
    "額外下載 Aaron Judge、Mookie Betts、Ronald Acuña Jr. 三位代表性球員的 "
    "2024 年逐球資料（各約 700–900 筆），用於熱區圖視覺化分析。")

# ══════════════════════════════════════════════════════════
# §3 資料處理流程
# ══════════════════════════════════════════════════════════
add_heading(doc, "3. 資料處理流程", 1)

add_heading(doc, "3.1 整體管線", 2)
add_figure(doc,
    POST_DIR / "fig_pipeline.png",
    "圖 1：從 Statcast 逐球資料到最終 wRC+ 預測的完整資料管線。",
    badge="🏆 海報重點圖  ｜  📑 期末簡報必放",
    width_cm=15)

add_heading(doc, "3.2 逐季打者面板建立", 2)
add_para(doc, "程式：bb_pipeline/statcast_season.py　輸出：data/raw/batting_by_season_2018_2024.parquet")
add_para(doc,
    "保留打席結束事件（PA_EVENTS），依 (batter, game_year) 分組計算以下指標：")
add_table(doc,
    headers=["計算指標", "公式／說明"],
    rows=[
        ["PA", "有效打席數（計入 woba_denom 之列）"],
        ["K%", "三振數 ÷ PA"],
        ["BB%", "四壞球數 ÷ PA"],
        ["BABIP", "安打（扣全壘打）÷ 場內擊球出局（BIP）"],
        ["BIP%", "場內擊球事件數 ÷ PA"],
        ["wOBA", "Σ(woba_value) ÷ Σ(woba_denom)"],
        ["wRC+", "以當季全聯盟 PA 加權 lg_wOBA，套用簡化公式（wOBA_scale=1.18）"],
        ["age_bat_median", "該季所有打席的年齡中位數"],
    ],
    col_widths=[4, 13.5]
)

add_heading(doc, "3.3 Lag 特徵工程", 2)
add_para(doc, "程式：bb_pipeline/projection_dataset.py　輸出：data/processed/projection_panel.parquet（943 列）")
add_para(doc,
    "對每位打者、每個目標球季 S，組成 lag 特徵面板。"
    "必要條件：S-1、S-2、S-3 各季 PA ≥ 100。")
add_table(doc,
    headers=["特徵類型", "欄位", "說明"],
    rows=[
        ["標籤", "wRCplus_target", "目標年 S 的 wRC+"],
        ["比率特徵（12欄）", "K%, BB%, BABIP, BIP% × lag1/2/3", "前三季各指標"],
        ["打席數（4欄）", "PA_lag1/2/3, PA_sum_lag3", "各季打席與三年總計"],
        ["年齡", "age_t", "取自 S-1 季年齡中位數"],
        ["守備位置", "primary_pos", "Lahman 資料；失敗時為 UNK"],
    ],
    col_widths=[4, 6, 7.5]
)

add_figure(doc,
    POST_DIR / "fig_timesplit.png",
    "圖 2：時間切分策略。Train（目標年 2021–22，n=453）、Validation（2023，n=234）、"
    "Test（2024，n=256）。箭頭示意預測 2024 時所需的三年 lag 資料範圍，"
    "此設計確保訓練時看不到未來資料。",
    badge="🏆 海報重點圖  ｜  📑 期末簡報必放",
    width_cm=15)

# ══════════════════════════════════════════════════════════
# §4 探索性資料分析
# ══════════════════════════════════════════════════════════
add_heading(doc, "4. 探索性資料分析（EDA）", 1)

add_heading(doc, "4.1 主要打擊指標分布", 2)
add_figure(doc,
    FIG_DIR / "eda_distributions.png",
    "圖 3：2018–2024 MLB 合格打者（PA ≥ 150，n=958）六項主要指標分布直方圖。"
    "wRC+ 接近常態分布，中心約 90；WAR 右偏，多數球員貢獻有限；K% 平均約 22%。",
    badge="📑 期末簡報用",
    width_cm=15)

add_para(doc, "關鍵統計：", bold=True)
add_table(doc,
    headers=["指標", "均值", "觀察"],
    rows=[
        ["wRC+", "≈ 90", "略低於聯盟平均 100，因弱打者也入列"],
        ["WAR", "≈ 4（中位數 ≈ 2）", "高度右偏，頂級球員 WAR 可達 40+"],
        ["K%", "≈ 22%", "現代棒球三振率持續上升的反映"],
        ["BB%", "≈ 8%", "選球型打者可達 15%+"],
    ],
    col_widths=[3.5, 4, 10]
)

add_heading(doc, "4.2 指標相關係數分析", 2)
add_figure(doc,
    FIG_DIR / "eda_correlation.png",
    "圖 4：九項核心打擊指標的 Pearson 相關係數矩陣。深紅代表強正相關，藍色代表負相關。",
    badge="📑 期末簡報用",
    width_cm=13)

add_para(doc, "關鍵發現：", bold=True)
add_table(doc,
    headers=["指標對", "相關係數", "解讀"],
    rows=[
        ["wRC+ ↔ OBP", "+0.86", "上壘率最能解釋打擊貢獻，選球與不出局是核心"],
        ["wRC+ ↔ SLG", "+0.85", "長打率次之，長打力直接驅動得分能力"],
        ["SLG ↔ ISO", "+0.89", "孤立長打率幾乎由 SLG 決定"],
        ["AVG ↔ K%", "−0.62", "高三振率球員打擊率普遍偏低"],
        ["WAR ↔ wRC+", "+0.65", "攻擊是 WAR 主成分，但守備貢獻也有影響"],
        ["BB% ↔ AVG", "−0.02", "選球能力與打擊率幾乎無關（如 Juan Soto）"],
    ],
    col_widths=[3.5, 3, 11]
)

add_heading(doc, "4.3 wRC+ 排名分析", 2)
add_figure(doc,
    FIG_DIR / "fig1_wrcplus_ranking.png",
    "圖 5：2018–2024 所有合格球季中 wRC+ 最高的前 20 名打者（依 K-Means 分群上色）。"
    "虛線為聯盟平均（100）。前 2 名（Nathan Lukes 186、Max Schrock 174）因打席數極少，"
    "wRC+ 受小樣本放大，不如 Mike Trout（3,600+ PA）和 Aaron Judge（3,700+ PA）具代表性。",
    badge="📑 期末簡報用",
    width_cm=13)

add_heading(doc, "4.4 WAR vs wRC+ 四象限分析", 2)
add_figure(doc,
    FIG_DIR / "fig2_war_vs_wrcplus.png",
    "圖 6：打者 WAR 對 wRC+ 的四象限分布圖。右上為「明星球員」（Aaron Judge、Juan Soto 等），"
    "左上為「守備價值型」，右下為「攻擊潛力型」（守備較弱），左下為「替補水準」。",
    badge="🏆 海報重點圖  ｜  📑 期末簡報必放",
    width_cm=14)

# ══════════════════════════════════════════════════════════
# §5 Statcast 熱區視覺化
# ══════════════════════════════════════════════════════════
add_heading(doc, "5. Statcast 熱區視覺化", 1)

add_heading(doc, "5.1 三球員進壘點熱區對比", 2)
add_figure(doc,
    FIG_DIR / "fig4_heatmap_all.png",
    "圖 7：Aaron Judge、Mookie Betts、Ronald Acuña Jr. 在 2024 球季所有打席的進壘點密度分布。"
    "黑框為標準好球帶。Judge 集中於中低區，Betts 分布均勻，Acuña 外側追打球較多。",
    badge="🏆 海報重點圖  ｜  📑 期末簡報必放",
    width_cm=15)

add_heading(doc, "5.2 依打擊結果分色的進壘點分布", 2)
add_figure(doc,
    FIG_DIR / "fig5_heatmap_by_result.png",
    "圖 8：三球員合計進壘點，依打擊結果分色。全壘打（紅星）集中於好球帶中央偏下，"
    "三振（紅 ×）分布廣泛，外角低球為高三振率區域。",
    badge="📑 期末簡報用",
    width_cm=10)

add_heading(doc, "5.3 Aaron Judge 強弱區對比", 2)
add_figure(doc,
    FIG_DIR / "fig10_judge_power_vs_weak.png",
    "圖 9：Aaron Judge 2024 球季強區（HR+2B+3B，n=94）與弱區（三振 K，n=164）對比。"
    "長打集中於好球帶中央，三振集中於外角右側。即使是頂級強打者，外角仍是弱點。",
    badge="📑 期末簡報用",
    width_cm=13)

add_heading(doc, "5.4 實際打擊率 vs 期望打擊率（運氣修正）", 2)
add_figure(doc,
    FIG_DIR / "fig9_xba_vs_ba.png",
    "圖 10：2024 球季打者實際 BA 對期望 BA（xBA）。對角線上方為「打得比應有水準好」，"
    "下方為「運氣較差」或受強守備壓制。顏色代表 BA−xBA 差值。",
    badge="📑 期末簡報用（補充）",
    width_cm=12)

# ══════════════════════════════════════════════════════════
# §6 球員分群分析
# ══════════════════════════════════════════════════════════
add_heading(doc, "6. 球員分群分析", 1)

add_heading(doc, "6.1 K-Means 分群（4 種打者類型）", 2)
add_para(doc,
    "以 AVG、OBP、SLG、BB%、接觸率（contact_rate）、孤立長打率（ISO）六項標準化指標，"
    "對 958 名合格打者進行 K-Means 聚類（k=4）。")
add_table(doc,
    headers=["分群", "平均 wRC+", "特徵描述", "典型代表"],
    rows=[
        ["強打型", "117.0", "高 OBP、高 SLG，低 K%（相對）", "Aaron Judge, Juan Soto"],
        ["選球型", "93.1", "高 BB%，SLG 中等，耐心選球", "耐心型打者"],
        ["接觸型", "90.6", "低 K%，高 AVG，但長打力有限", "接觸型打者"],
        ["工具型", "56.4", "低 AVG、高 K%，各指標偏弱", "多為備用球員"],
    ],
    col_widths=[3, 2.5, 7, 5]
)

add_heading(doc, "6.2 明星球員雷達圖", 2)
add_figure(doc,
    POST_DIR / "fig_radar_stars.png",
    "圖 11：Aaron Judge、Juan Soto、Shohei Ohtani、Mookie Betts、Freddie Freeman "
    "六項打擊指標的標準化能力比較（各指標取各球員生涯最優值，標準化至 0–1）。"
    "Judge 長打力最強，Soto 選球最佳，Ohtani 最均衡。",
    badge="📑 期末簡報用（補充）",
    width_cm=11)

# ══════════════════════════════════════════════════════════
# §7 隔年 wRC+ 預測模型
# ══════════════════════════════════════════════════════════
add_heading(doc, "7. 隔年 wRC+ 預測模型", 1)

add_heading(doc, "7.1 模型設計", 2)
add_table(doc,
    headers=["項目", "說明"],
    rows=[
        ["任務", "回歸問題：輸入前三季 lag 特徵，輸出下一季 wRC+（連續值）"],
        ["主要模型", "Random Forest Regressor（n_estimators=400, max_depth=12, min_samples_leaf=8）"],
        ["備選模型", "XGBoost（macOS 需 brew install libomp）"],
        ["特徵預處理", "數值特徵直通；primary_pos 經 OneHotEncoder（handle_unknown=ignore）"],
        ["Train", "target_season ∈ {2021, 2022}，共 453 筆"],
        ["Validation", "target_season = 2023，共 234 筆（Permutation Importance 計算用）"],
        ["Test", "target_season = 2024，共 241 筆（過濾異常值後）"],
    ],
    col_widths=[3.5, 14]
)

add_heading(doc, "7.2 特徵重要性分析", 2)
add_figure(doc,
    POST_DIR / "fig_perm_importance.png",
    "圖 12：Random Forest 在 Validation（2023）上的 Permutation Importance。"
    "橫軸為打亂該特徵後驗證誤差上升量，越長代表模型越依賴該特徵。"
    "紅色為關鍵特徵（Top 30%）。",
    badge="🏆 海報重點圖  ｜  📑 期末簡報必放",
    width_cm=14)

add_para(doc, "關鍵發現：", bold=True)
add_table(doc,
    headers=["排名", "特徵", "解讀"],
    rows=[
        ["1", "三年 PA 總計", "穩定出賽是健康度信號，代表球員維持競爭力的能力"],
        ["2", "四壞% (3yr前)", "選球能力長期穩定，早年 BB% 仍能預測未來"],
        ["3", "三振% (1yr前)", "最近一年三振率對下季表現影響顯著"],
        ["4", "打席數 PA (1yr前)", "近年出賽量預示下季參與度"],
        ["5", "年齡", "生涯曲線效應（27–29 歲巔峰）確實存在"],
        ["負值", "BABIP (3yr前)", "遠期 BABIP 負重要性，確認其均值回歸（噪音）特性"],
    ],
    col_widths=[1.5, 4.5, 11.5]
)

add_warning_box(doc,
    "注意：Permutation Importance 衡量模型依賴程度，非因果關係。"
    "BB% 排前列不代表四壞率能「提升」未來 wRC+，而是兩者在時間上共變化。")

add_heading(doc, "7.3 預測結果（Hold-out Test 2024）", 2)
add_figure(doc,
    POST_DIR / "fig_actual_vs_pred.png",
    "圖 13：Hold-out Test 2024（n=241，已過濾 wRC+<20 異常值）實際 vs RF 預測散點圖。"
    "虛線為完美預測線（y=x）。顏色代表預測誤差大小（紅色誤差大）。"
    "Juan Soto、Shohei Ohtani 等超級球星被低估；低 wRC+ 球員被高估（向均值壓縮）。",
    badge="🏆 海報重點圖  ｜  📑 期末簡報必放",
    width_cm=12)

add_table(doc,
    headers=["評估指標", "數值", "解讀"],
    rows=[
        ["MAE", "22.1 wRC+ 點", "wRC+ 年際自然波動也約 15–25 點，此誤差接近自然波動上限"],
        ["R²", "−0.018", "負值：系統性偏差（向均值壓縮）主導，略遜於均值預測"],
    ],
    col_widths=[3.5, 3.5, 10.5]
)

# ══════════════════════════════════════════════════════════
# §8 分析結論
# ══════════════════════════════════════════════════════════
add_heading(doc, "8. 分析結論", 1)

add_heading(doc, "8.1 主要發現", 2)

add_para(doc, "EDA 層面：", bold=True)
for item in [
    "OBP 和 SLG 與 wRC+ 高度相關（r > 0.85），是球員攻擊價值的核心指標",
    "BB%（選球能力）與 AVG 近零相關，說明高選球型打者未必打擊率高，但仍能創造高 wRC+",
    "進階指標（wRC+）能識別被 AVG 低估的球員（如 Jacob Nottingham：AVG .186，wRC+ 122.7）",
    "熱區分析確認現代棒球投手策略：外角低球是三振熱區，中央偏下是長打爆發區",
]:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(item)
    run.font.size = Pt(10.5)

add_para(doc, "預測模型層面：", bold=True)
for item in [
    "三年 PA 總計是最重要預測特徵：健康穩定出賽是隔年表現的最強信號",
    "四壞率（BB%）具跨年穩定性：三年前的 BB% 仍能預測未來，反映打者眼力是長期能力",
    "BABIP 均值回歸確認：遠期 BABIP 呈負重要性，不應過度解讀單年 BABIP 異常",
    "年齡效應可測量：生涯曲線對 wRC+ 有貢獻，27–29 歲前後出現系統性差異",
    "Test MAE = 22.1 wRC+ 點，接近 wRC+ 年際自然波動；R² = −0.018 誠實反映模型侷限",
]:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(item)
    run.font.size = Pt(10.5)

add_heading(doc, "8.2 模型限制", 2)
add_table(doc,
    headers=["限制", "說明"],
    rows=[
        ["訓練資料量有限", "有監督 target_season 僅 4 年（2021–24），Train 只有 2 個目標年"],
        ["wRC+ 簡化公式", "非 FanGraphs 官方公式，可能有 1–3 點差異"],
        ["守備位置品質", "Lahman 資料部分缺失，位置特徵多為 UNK，降低位置的預測力"],
        ["未含受傷資訊", "球員因傷缺賽會影響 PA，模型無法區分技術衰退與受傷"],
        ["均值壓縮效應", "RF 預測天然向中間壓縮，對超級球星低估、弱打者高估"],
    ],
    col_widths=[4.5, 13]
)

add_heading(doc, "8.3 未來改進方向", 2)
for i, item in enumerate([
    "擴增訓練年份：等待 2025 球季結束後延伸至 target_season=2025",
    "納入擊球品質指標：launch_speed、launch_angle、Hard Hit% 等 Statcast 物理量",
    "XGBoost 系統性比較：安裝 libomp 後與 RF 並列評估",
    "Walk-Forward Validation：逐年擴張訓練窗，更準確估計真實泛化能力",
    "位置分層建模：分捕手、外野手、內野手建立子模型，降低 wRC+ 基準線差異的干擾",
], 1):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(item)
    run.font.size = Pt(10.5)

# ══════════════════════════════════════════════════════════
# 附錄：圖片使用建議
# ══════════════════════════════════════════════════════════
add_heading(doc, "附錄：圖片使用建議", 1)

add_heading(doc, "海報（Poster）重點圖 — 建議放 5–6 張", 2)
add_table(doc,
    headers=["圖", "檔案", "說明"],
    rows=[
        ["圖 1 管線流程", "poster/fig_pipeline.png", "讓觀眾快速理解整體研究設計"],
        ["圖 2 時間切分", "poster/fig_timesplit.png", "展示嚴謹的時間切分設計"],
        ["圖 6 WAR vs wRC+", "fig2_war_vs_wrcplus.png", "直觀展示球員價值分布與分群"],
        ["圖 7 熱區圖", "fig4_heatmap_all.png", "視覺衝擊力強，吸引非棒球觀眾"],
        ["圖 12 特徵重要性", "poster/fig_perm_importance.png", "展示模型核心發現"],
        ["圖 13 實際 vs 預測", "poster/fig_actual_vs_pred.png", "呈現模型表現與限制"],
    ],
    col_widths=[3, 5.5, 9]
)

add_heading(doc, "不建議使用（舊同年模型，易造成誤解）", 2)
add_warning_box(doc,
    "fig6_model_comparison.png（同年 5-fold CV，R²≈0.86 嚴重誤導）、"
    "fig7_feature_importance.png（同年 MDI，OBP/SLG 主導）、"
    "fig8_actual_vs_predicted.png（同年 scatter，非 hold-out）"
    "——這三張是舊的同年預測模型圖，容易讓評審以為存在 data leakage，請勿放入報告或海報。")

# ══════════════════════════════════════════════════════════
# 儲存
# ══════════════════════════════════════════════════════════
doc.save(str(OUT))
print(f"Saved → {OUT}")
